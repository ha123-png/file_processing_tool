import os
from pathlib import Path
from app.shared import get_config_snapshot

def detect_encoding(filepath):
    with open(filepath, "rb") as f:
        raw = f.read()
    encodings = ["utf-8", "gbk", "gb2312", "gb18030", "utf-16", "big5", "shift_jis", "euc-kr"]
    for enc in encodings:
        try:
            raw.decode(enc)
            return enc
        except (UnicodeDecodeError, LookupError):
            continue
    import locale
    return locale.getpreferredencoding()

def read_text_file(filepath):
    encoding = detect_encoding(filepath)
    with open(filepath, "r", encoding=encoding, errors="replace") as f:
        return f.read()

def read_docx_file(filepath):
    from docx import Document
    doc = Document(filepath)
    lines = [p.text for p in doc.paragraphs]
    return "\n".join(lines)

def _docx_has_images(filepath):
    from docx import Document
    doc = Document(filepath)
    for rel in doc.part.rels.values():
        if "image" in rel.reltype:
            return True
    return False

def _extract_docx_images(doc, output_prefix):
    image_paths = []
    for i, rel in enumerate(doc.part.rels.values()):
        if "image" in rel.reltype:
            try:
                image_part = rel.target_part
                ext = image_part.partname.split(".")[-1] if "." in image_part.partname else "png"
                if ext.lower() not in ("png", "jpg", "jpeg", "bmp", "gif", "webp"):
                    ext = "png"
                img_path = str(output_prefix) + f"_docx_img{i}.{ext}"
                with open(img_path, "wb") as f:
                    f.write(image_part.blob)
                image_paths.append(img_path)
            except Exception:
                pass
    return image_paths

def read_docx_with_images(filepath):
    from docx import Document
    doc = Document(filepath)
    paragraphs = [p.text.strip() for p in doc.paragraphs]
    para_text = "\n".join([t for t in paragraphs if t])

    image_infos = []
    for i, p in enumerate(doc.paragraphs):
        has_image = False
        for run in p.runs:
            for inline in run._element.findall('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}drawing'):
                has_image = True
                break
            for inline in run._element.findall('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}pict'):
                has_image = True
                break
        if has_image:
            before = ""
            after = ""
            for j in range(i - 1, -1, -1):
                pt = paragraphs[j]
                if pt:
                    before = pt
                    break
            for j in range(i + 1, len(paragraphs)):
                pt = paragraphs[j]
                if pt:
                    after = pt
                    break
            image_infos.append({"para_index": i, "before": before, "after": after})

    image_paths = _extract_docx_images(doc, filepath)
    for idx, info in enumerate(image_infos):
        if idx < len(image_paths):
            info["path"] = image_paths[idx]

    return para_text, image_infos
